# C:\F-ChatAI\ingest_windows.py
import unicodedata, hashlib, re
import pathlib
import psycopg2
import json
import os
def _pick_threads():
    # ручной оверрайд через окружение
    v = os.getenv("INGEST_CPU_THREADS")
    if v:
        return max(1, int(v))
    # эвристика под Ryzen AI 9 365: половина логических потоков, но в [6..8]
    try:
        cores_logical = os.cpu_count() or 8
    except Exception:
        cores_logical = 8
    return max(6, min(8, cores_logical // 2))

THREADS = _pick_threads()

# единый лимит для BLAS/OMP/Rayon
for k in (
    "OMP_NUM_THREADS", "MKL_NUM_THREADS", "OPENBLAS_NUM_THREADS",
    "NUMEXPR_NUM_THREADS", "BLIS_NUM_THREADS", "RAYON_NUM_THREADS"
):
    os.environ.setdefault(k, str(THREADS))

# фиксируем поведение OpenMP/oneDNN на CPU (AMD)
os.environ.setdefault("OMP_DYNAMIC", "FALSE")
os.environ.setdefault("OMP_WAIT_POLICY", "PASSIVE")   # можно попробовать "ACTIVE" для максимальной скорости
os.environ.setdefault("KMP_AFFINITY", "granularity=fine,compact,1,0")

# токенайзеры HF (rust/rayon)
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
import sys
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from sentence_transformers import SentenceTransformer
from tqdm import tqdm
import logging
from collections import Counter
import numpy as np
try:
    import torch
    torch.set_num_threads(THREADS)
    torch.set_num_interop_threads(1)
except Exception:
    pass
import mimetypes
try:
    import docx  # from python-docx
except Exception:
    docx = None
import zipfile
import datetime

try:
    import openpyxl
    from openpyxl import load_workbook
    from openpyxl.utils import range_boundaries
except Exception:
    openpyxl = None

try:
    from transformers import AutoTokenizer
except Exception:
    AutoTokenizer = None

# PDF: опциональные зависимости
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import pdfplumber  # на базе pdfminer.six — пригодится для таблиц/фолбэка
except Exception:
    pdfplumber = None

NS = {
    "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r":   "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v":   "urn:schemas-microsoft-com:vml",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}

SPACE_MAP = {0x00A0:' ',0x1680:' ',0x2000:' ',0x2001:' ',0x2002:' ',0x2003:' ',0x2004:' ',
             0x2005:' ',0x2006:' ',0x2007:' ',0x2008:' ',0x2009:' ',0x200A:' ',0x202F:' ',
             0x205F:' ',0x3000:' '}
INVISIBLES = {0x200B,0x200C,0x200D,0xFEFF,0x00AD}

# Исправление кодировки для Windows
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
    os.environ['PGCLIENTENCODING'] = 'UTF8'

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('ingest.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)
MEDIA_ROOT = pathlib.Path(os.getenv("MEDIA_ROOT", r"C:\F-ChatAI\media"))

@dataclass
class Document:
    """Структура документа для обработки"""
    book: str
    section: str
    page: Optional[int]
    text: str
    meta: Dict = field(default_factory=dict)
    doc_hash: Optional[str] = None
    parent_section: Optional[str] = None
    subsection: Optional[str] = None
    importance_score: float = 1.0
    is_parent: bool = False
    parent_group: Optional[str] = None
    child_index: Optional[int] = None
    parent_title: Optional[str] = None

class AdvancedDocumentProcessor:
    """Улучшенный процессор документов с семантическим анализом"""
    
    def __init__(self, db_config: dict, model_name: str = "BAAI/bge-m3"):
        self.db_config = self._fix_db_config(db_config)
        self.model = SentenceTransformer(model_name)
        self.emb_dim = getattr(self.model, "get_sentence_embedding_dimension", lambda: None)() or 1024
        self.tokenizer = None
        if AutoTokenizer is not None:
            try:
                # для bge-m3 подходит его же токенизатор
                self.tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True)
            except Exception:
                # запасной вариант, если модель токенизатора не подтянулась
                try:
                    self.tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-m3", use_fast=True)
                except Exception:
                    self.tokenizer = None

        if self.tokenizer is not None:
            try:
                # очень большой порог, чтобы HF не предупреждал
                self.tokenizer.model_max_length = int(1e12)
                # на новых версиях можно пометить предупреждение как уже показанное
                if hasattr(self.tokenizer, "deprecation_warnings"):
                    self.tokenizer.deprecation_warnings[
                        "sequence_length_is_longer_than_model_max_length"
                    ] = True
            except Exception:
                pass

        # размеры в ТОКЕНАХ (а не в символах)
        self.chunk_tokens = 450         # целевой размер чанка
        self.overlap_tokens = 64         # перекрытие между чанками
        self.min_chunk_tokens = 120      # слишком мелкие куски стараемся присоединять
        self.max_chunk_tokens = 1200     # «очень длинная секция» → принудительно дробим

        # символьные пороги оставляем как бэкап (используются при отсутствии tokenizer)
        self.chunk_size = 800  # Оптимальный размер для контекста
        self.overlap = 100     # Перекрытие для сохранения контекста
        self.min_chunk_size = 200
        self.max_chunk_size = 1500

        self.overlap_tokens = min(self.overlap_tokens, max(0, self.chunk_tokens // 2))

        # Паттерны для определения структуры документа
        self.header_patterns = [
            (r'^#{1,3}\s+(.+)$', 3),  # Markdown заголовки
            (r'^(\d+\.?\d*)\s+([А-ЯA-Z].+)$', 2),  # Нумерованные заголовки
            (r'^([А-ЯA-Z][^.!?]{3,50}):?\s*$', 1),  # Короткие заголовки
            (r'^Глава\s+(\d+|[IVXLCDM]+)[.:]?\s*(.+)?$', 3),  # Главы
            (r'^Раздел\s+(\d+|[IVXLCDM]+)[.:]?\s*(.+)?$', 2),  # Разделы
            (r'^Приложение\s+([А-ЯA-Z\d]+)[.:]?\s*(.+)?$', 2),  # Приложения
        ]
        
        # Ключевые термины для определения важности
        self.important_terms = {
            'важно': 3, 'внимание': 3, 'обязательно': 3, 'критично': 3,
            'примечание': 2, 'замечание': 2, 'рекомендация': 2,
            'пример': 1, 'дополнительно': 1, 'справка': 1,
            'настройка': 2, 'конфигурация': 2, 'параметр': 2,
            'ошибка': 3, 'проблема': 3, 'решение': 3,
            'алгоритм': 2, 'процедура': 2, 'функция': 2
        }

    def _tok_len(self, s: str) -> int:
        """Длина текста в токенах; если токенизатора нет — грубая эвристика по символам."""
        if not self.tokenizer:
            # эвристика: ~4 символа на токен для кириллицы/латиницы вперемешку
            return max(1, len(s) // 4)
        return len(self.tokenizer.encode(s, add_special_tokens=False))

            
    def _fix_db_config(self, config: dict) -> dict:
        """Исправление конфигурации БД для Windows"""
        fixed_config = config.copy()
        fixed_config['options'] = '-c client_encoding=UTF8'
        if fixed_config.get('host') == 'localhost':
            fixed_config['host'] = '127.0.0.1'
        return fixed_config
    
    @staticmethod
    def _media_key_for(filepath: pathlib.Path) -> str:
        # стабильный ключ каталога для картинок конкретного исходного файла
        base = str(filepath.resolve()).encode("utf-8", "ignore")
        return hashlib.md5(base).hexdigest()[:12]
    
    def _read_csv(self, path: pathlib.Path) -> tuple[str, list[dict]]:
        """
        Чтение CSV -> markdown-таблицы с разбиением на блоки строк.
        Возвращает (markdown_text, images). Для CSV images всегда [].
        """
        import csv
 
        # лимиты
        MAX_COLS = int(os.getenv("INGEST_CSV_MAX_COLS", 60))
        MAX_ROWS = int(os.getenv("INGEST_CSV_MAX_ROWS", 20000))
        ROWS_PER_BLOCK = int(os.getenv("INGEST_CSV_ROWS_PER_BLOCK", 1000))
        MAX_CELL = int(os.getenv("INGEST_CSV_MAX_CELL_CHARS", 2000))

        def esc_cell(s: str) -> str:
            if s is None:
                return ""
            s = str(s).replace("|", "\\|").replace("\r\n", " ").replace("\r", " ").replace("\n", " ").strip()
            if len(s) > MAX_CELL:
                s = s[:MAX_CELL].rstrip() + "…"
            return s

        # определим кодировку и подготовим reader
        enc = self._detect_encoding(path)
        if enc == "unknown":
            enc = "utf-8"

        text_lines: list[str] = [f"## CSV: {path.name}"]
        images: list[dict] = []

        with open(path, "r", encoding=enc, newline="") as f:
            sample = f.read(128 * 1024)
            f.seek(0)

            # диалект (делимитеры: , ; tab |)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=[",", ";", "\t", "|"])
            except Exception:
                # простая эвристика
                if sample.count(";") > sample.count(",") and sample.count(";") >= sample.count("\t"):
                    d = csv.excel()
                    d.delimiter = ";"
                    dialect = d
                elif sample.count("\t") >= max(sample.count(","), sample.count(";")):
                    d = csv.excel_tab()
                    dialect = d
                elif sample.count("|") > 0 and sample.count("|") >= max(sample.count(","), sample.count(";"), sample.count("\t")):
                    d = csv.excel()
                    d.delimiter = "|"
                    dialect = d
                else:
                    dialect = csv.excel()

            try:
                has_header = csv.Sniffer().has_header(sample)
            except Exception:
                has_header = True

            reader = csv.reader(f, dialect)

            # прочитаем первую строку (возможный header)
            first_row = next(reader, None)
            if first_row is None:
                return "\n".join(text_lines), images

            if has_header:
                header = [esc_cell(x) for x in first_row[:MAX_COLS]]
            else:
                # если хедера нет — сгенерируем
                ncols = min(len(first_row), MAX_COLS)
                header = [f"col{i}" for i in range(1, ncols + 1)]

            # печатаем блоками
            def start_block(block_start_idx: int):
                text_lines.append(f"### Строки {block_start_idx}-{min(block_start_idx + ROWS_PER_BLOCK - 1, MAX_ROWS)}")
                text_lines.append(" | ".join(header))
                text_lines.append(" | ".join(["---"] * len(header)))

            # если хедера не было — первая строка реально данные
            from itertools import chain
            data_iter = reader if has_header else chain([first_row], reader)

            written = 0
            block_row_start = 1
            start_block(block_row_start)
            block_count = 0

            for row in data_iter:
                if written >= MAX_ROWS:
                    break
                row = [esc_cell(x) for x in row[:len(header)]]
                text_lines.append(" | ".join(row))
                written += 1
                block_count += 1

                if block_count >= ROWS_PER_BLOCK and written < MAX_ROWS:
                    block_row_start = written + 1
                    block_count = 0
                    start_block(block_row_start)

            # если файл имел больше колонок — подсказка
            if len(first_row) > len(header):
                text_lines.append("")
                text_lines.append(f"_Примечание: отображены только первые {len(header)} колонок из {len(first_row)}._")

            if written == 0:
                # табличка хотя бы с заголовком (пустая)
                text_lines.append("_Нет данных (после заголовка)._")

        return "\n".join(text_lines), images

    
    def _read_pdf(self, path: pathlib.Path) -> tuple[str, list[dict]]:
        """
        Чтение PDF: извлекаем текст постранично и вытаскиваем изображения.
        Возвращает (markdown_text, images), где images — список словарей
        {"rel": "...", "mime": "image/png", "sha256": "...", "page": int}.
        """
        media_key = self._media_key_for(path)
        media_root = MEDIA_ROOT
        media_root.mkdir(parents=True, exist_ok=True)

        images: list[dict] = []
        text_lines: list[str] = []

        def _save_pixmap(pix: "fitz.Pixmap", page_no: int, idx: int) -> str:
            # Преобразуем в RGB при необходимости и сохраняем PNG
            if pix.alpha or pix.n > 4:  # CMYK/с альфой -> RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)

            rel = f"{media_key}/p{page_no:04d}_img_{idx:03d}.png"
            full = media_root / rel
            full.parent.mkdir(parents=True, exist_ok=True)
            pix.save(str(full))
            with open(full, "rb") as f:
                sha = hashlib.sha256(f.read()).hexdigest()
            images.append({
                "rel": rel.replace("\\", "/"),
                "mime": "image/png",
                "sha256": sha,
                "page": page_no
            })
            return rel.replace("\\", "/")

        # --- основной путь: PyMuPDF ---
        if fitz is not None:
            try:
                doc = fitz.open(str(path))
                if getattr(doc, "needs_pass", False):
                    # пробуем пустой пароль; иначе — исключение
                    if not doc.authenticate(""):
                        raise RuntimeError("PDF защищён паролем")

                for pno in range(len(doc)):
                    page = doc[pno]
                    text_lines.append(f"## Страница {pno + 1}")

                    # Текст страницы (сохраняем базовую верстку)
                    try:
                        txt = page.get_text(
                            "text",
                            flags=getattr(fitz, "TEXT_PRESERVE_LIGATURES", 0)
                                | getattr(fitz, "TEXT_PRESERVE_WHITESPACE", 0)
                        ).strip()
                    except Exception:
                        txt = page.get_text("text").strip()
                    text_lines.append(txt if txt else "")

                    # Изображения страницы
                    imgs = page.get_images(full=True)
                    for idx, (xref, *_rest) in enumerate(imgs, start=1):
                        try:
                            pix = fitz.Pixmap(doc, xref)
                            rel = _save_pixmap(pix, pno + 1, idx)
                            # Вставим превью в текст (удобно для RAG навигации)
                            text_lines.append(f"![p{pno + 1:04d}_img_{idx:03d}](/media/{rel})")
                        except Exception:
                            continue

                    text_lines.append("")  # пустая строка-разделитель
                doc.close()
                return "\n".join(text_lines).strip(), images
            except Exception as e:
                # Падаем на фолбэк ниже
                logger.warning(f"PyMuPDF не справился с {path.name}: {e}")

        # --- фолбэк: только текст через pdfplumber ---
        if pdfplumber is not None:
            try:
                with pdfplumber.open(str(path)) as pdf:
                    for pno, page in enumerate(pdf.pages, start=1):
                        text_lines.append(f"## Страница {pno}")
                        txt = (page.extract_text() or "").strip()
                        text_lines.append(txt)
                        # Таблицы (best-effort) -> markdown
                        try:
                            tables = page.extract_tables() or []
                            for t in tables:
                                if not t or not any(any(c) for c in t):
                                    continue
                                # первая строка — заголовок
                                header = [ (c or "").strip() for c in t[0] ]
                                text_lines.append(" | ".join(header))
                                text_lines.append(" | ".join(["---"] * len(header)))
                                for row in t[1:]:
                                    row = [ (c or "").strip() for c in row ]
                                    text_lines.append(" | ".join(row))
                                text_lines.append("")
                        except Exception:
                            pass
                        text_lines.append("")
                return "\n".join(text_lines).strip(), images
            except Exception as e:
                logger.error(f"pdfplumber не справился с {path.name}: {e}")

        # Совсем без зависимостей — ничего сделать нельзя
        raise RuntimeError(
            "Для PDF нужен хотя бы один из пакетов: 'pymupdf' (рекомендуется) или 'pdfplumber'."
        )


    def _extract_pdf_metadata(self, path: pathlib.Path) -> dict:
        """
        Best-effort метаданные PDF (title, author, creation/mod dates, pages).
        Не критично, если библиотек нет — вернём пусто.
        """
        meta: dict = {}
        # PyMuPDF быстрее и даёт больше полей
        if fitz is not None:
            try:
                doc = fitz.open(str(path))
                md = doc.metadata or {}
                if md:
                    meta["pdf_meta"] = {
                        "title": md.get("title"),
                        "author": md.get("author"),
                        "subject": md.get("subject"),
                        "keywords": md.get("keywords"),
                        "creationDate": md.get("creationDate"),
                        "modDate": md.get("modDate"),
                        "producer": md.get("producer"),
                        "creator": md.get("creator"),
                        "pages": len(doc)
                    }
                doc.close()
                return meta
            except Exception:
                pass

        # Фолбэк: pdfplumber/pdfminer
        if pdfplumber is not None:
            try:
                with pdfplumber.open(str(path)) as pdf:
                    md = getattr(pdf, "metadata", None) or {}
                    if md:
                        meta["pdf_meta"] = {
                            "title": md.get("Title"),
                            "author": md.get("Author"),
                            "subject": md.get("Subject"),
                            "keywords": md.get("Keywords"),
                            "creationDate": md.get("CreationDate"),
                            "modDate": md.get("ModDate"),
                            "producer": md.get("Producer"),
                            "creator": md.get("Creator"),
                            "pages": len(pdf.pages)
                        }
                return meta
            except Exception:
                pass
        return meta


    def _read_xlsx(self, path: pathlib.Path) -> tuple[str, list[dict]]:
        """
        Чтение Excel .xlsx -> markdown-таблицы + извлечение изображений.
        Улучшения: merged cells, гиперссылки, формат дат/чисел, галерея по листам.
        """
        if openpyxl is None:
            raise RuntimeError("openpyxl не установлен. Установите пакет: pip install openpyxl")

        import posixpath as pp
        from xml.etree import ElementTree as ET
        from openpyxl import load_workbook
        from openpyxl.utils import range_boundaries
        from openpyxl.utils.datetime import from_excel as excel_dt_from
        from openpyxl.styles.numbers import is_date_format

        # по умолчанию включаем галерею; можно отключить INGEST_XLSX_GALLERY=0
        WANT_GALLERY = os.getenv("INGEST_XLSX_GALLERY", "1") not in ("0", "false", "False", "no", "off")

        # --- извлечём медиа из xl/media/* и построим карту исходное_имя -> сохранённый путь ---
        media_key = self._media_key_for(path)
        media_root = MEDIA_ROOT
        media_root.mkdir(parents=True, exist_ok=True)
        images: list[dict] = []
        img_idx = 0
        media_name_to_rel: dict[str, str] = {}

        def _save_xlsx_image(bin_data: bytes, suggested_ext: str | None) -> str:
            nonlocal img_idx
            img_idx += 1
            ext = (suggested_ext or ".bin")
            rel = f"{media_key}/img_{img_idx:03d}{ext}"
            full = media_root / rel
            full.parent.mkdir(parents=True, exist_ok=True)
            with open(full, "wb") as f:
                f.write(bin_data)
            sha = hashlib.sha256(bin_data).hexdigest()
            mime = mimetypes.guess_type(str(full))[0] or "application/octet-stream"
            images.append({"rel": rel.replace("\\", "/"), "mime": mime, "sha256": sha})
            return rel.replace("\\", "/")

        try:
            with zipfile.ZipFile(str(path), "r") as zf:
                for info in zf.infolist():
                    if info.filename.lower().startswith("xl/media/") and not info.is_dir():
                        data = zf.read(info)
                        _, ext = os.path.splitext(info.filename)
                        rel = _save_xlsx_image(data, ext if ext else None)
                        media_name_to_rel[info.filename.replace("\\", "/")] = rel
        except Exception:
            pass  # best effort

        # --- сопоставим "лист -> картинки" через workbook/sheet/drawings ---
        sheet_to_media: dict[str, list[str]] = {}
        if WANT_GALLERY:
            try:
                with zipfile.ZipFile(str(path), "r") as zf:
                    # sheets: name -> r:id
                    wb_xml = ET.fromstring(zf.read("xl/workbook.xml"))
                    NS_MAIN = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
                    ns = {"r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
                    sheet_name_to_rid: dict[str, str] = {}
                    for sh in wb_xml.findall(f".//{{{NS_MAIN}}}sheet"):
                        name = sh.attrib.get("name")
                        rid = sh.attrib.get(f"{{{ns['r']}}}id")
                        if name and rid:
                            sheet_name_to_rid[name] = rid
                    # r:id -> worksheets/sheetN.xml
                    rels_xml = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
                    rid_to_target: dict[str, str] = {}
                    for rel in rels_xml.findall(".//Relationship"):
                        rid = rel.attrib.get("Id")
                        tgt = rel.attrib.get("Target")
                        typ = rel.attrib.get("Type", "")
                        if rid and tgt and typ.endswith("/worksheet"):
                            rid_to_target[rid] = tgt

                    for sheet_name, rid in sheet_name_to_rid.items():
                        sheet_part = rid_to_target.get(rid)
                        if not sheet_part:
                            continue
                        sheet_rels_path = f"xl/{pp.dirname(sheet_part)}/_rels/{pp.basename(sheet_part)}.rels"
                        if sheet_rels_path not in zf.namelist():
                            continue
                        srels_xml = ET.fromstring(zf.read(sheet_rels_path))
                        drawing_targets = []
                        for rel in srels_xml.findall(".//Relationship"):
                            typ = rel.attrib.get("Type", "")
                            if typ.endswith("/drawing"):
                                drawing_targets.append(rel.attrib.get("Target"))
                        media_for_sheet: list[str] = []
                        for d_target in drawing_targets:
                            drawing_xml_path = "xl/" + pp.normpath(pp.join(pp.dirname(sheet_part), d_target))
                            if drawing_xml_path not in zf.namelist():
                                continue
                            # rels для рисунка: rId -> xl/media/imageY.*
                            drawing_rels_path = f"{pp.dirname(drawing_xml_path)}/_rels/{pp.basename(drawing_xml_path)}.rels"
                            rid_to_media_target: dict[str, str] = {}
                            if drawing_rels_path in zf.namelist():
                                drels_xml = ET.fromstring(zf.read(drawing_rels_path))
                                for rel in drels_xml.findall(".//Relationship"):
                                    rid_d = rel.attrib.get("Id")
                                    tgt = rel.attrib.get("Target")
                                    typ = rel.attrib.get("Type", "")
                                    if rid_d and tgt and (typ.endswith("/image") or "/image" in typ):
                                        media_path = pp.normpath(pp.join(pp.dirname(drawing_xml_path), tgt))
                                        if not media_path.startswith("xl/"):
                                            media_path = "xl/" + media_path
                                        rid_to_media_target[rid_d] = media_path
                            # a:blip r:embed -> rId -> media
                            drawing_xml = ET.fromstring(zf.read(drawing_xml_path))
                            ns2 = {
                                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                                "xdr": "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing",
                                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                            }
                            for blip in drawing_xml.findall(".//xdr:blipFill/a:blip", ns2):
                                rid_embed = blip.attrib.get(f"{{{ns2['r']}}}embed")
                                if not rid_embed:
                                    continue
                                media_zip_path = rid_to_media_target.get(rid_embed)
                                if not media_zip_path:
                                    continue
                                rel_saved = media_name_to_rel.get(media_zip_path)
                                if rel_saved:
                                    media_for_sheet.append(rel_saved)
                        if media_for_sheet:
                            seen, uniq = set(), []
                            for rel_saved in media_for_sheet:
                                if rel_saved not in seen:
                                    seen.add(rel_saved)
                                    uniq.append(rel_saved)
                            sheet_to_media[sheet_name] = uniq
            except Exception:
                sheet_to_media = {}

        # --- конвертируем листы в markdown-таблицы ---
        wb = load_workbook(str(path), read_only=True, data_only=True)
        text_lines: list[str] = []

        MAX_COLS = int(os.getenv("INGEST_XLSX_MAX_COLS", 60))
        MAX_ROWS = int(os.getenv("INGEST_XLSX_MAX_ROWS", 5000))

        def esc_md_cell(s: str) -> str:
            return (s or "").replace("|", "\\|").replace("\n", " ").strip()

        def fmt_number(v: float) -> str:
            if abs(v - round(v)) < 1e-9 and abs(v) < 1e15:
                return str(int(round(v)))
            s = f"{v:.10f}".rstrip("0").rstrip(".")
            return s if s else "0"

        def fmt_cell(ws, cell) -> str:
            v = cell.value
            if v is None:
                # формула без кэша в data_only
                try:
                    if getattr(cell, "data_type", None) == "f":
                        return ""
                except Exception:
                    return ""
                return ""
            # даты/время: и объект, и числовой сериал
            try:
                if getattr(cell, "is_date", False):
                    if isinstance(v, (datetime.datetime, datetime.date, datetime.time)):
                        if isinstance(v, datetime.datetime):
                            return v.isoformat(sep=" ", timespec="seconds")
                        if isinstance(v, datetime.time):
                            return v.isoformat(timespec="seconds")
                        return v.isoformat()
                    if isinstance(v, (int, float)) and is_date_format(getattr(cell, "number_format", "")):
                        dt = excel_dt_from(v, wb.epoch)
                        if isinstance(dt, datetime.datetime):
                            return dt.isoformat(sep=" ", timespec="seconds")
                        return dt.isoformat()
            except Exception:
                pass
            if isinstance(v, float):
                return fmt_number(v)
            if isinstance(v, bool):
                return "TRUE" if v else "FALSE"
            return str(v)

        for ws in wb.worksheets:
            # используемую область
            try:
                dim = ws.calculate_dimension()
                min_col, min_row, max_col, max_row = range_boundaries(dim)
            except Exception:
                min_col, min_row, max_col, max_row = 1, 1, min(20, ws.max_column or 1), min(200, ws.max_row or 1)

            max_col = min(max_col, min_col + MAX_COLS - 1)
            max_row = min(max_row, min_row + MAX_ROWS - 1)

            # таблица merged: расстилаем top-left значение на всю merged область
            merged_areas: list[tuple[int, int, int, int, str]] = []
            try:
                for rng in getattr(ws, "merged_cells", []).ranges:
                    min_r, min_c, max_r, max_c = rng.min_row, rng.min_col, rng.max_row, rng.max_col
                    top_val = esc_md_cell(fmt_cell(ws, ws.cell(row=min_r, column=min_c)))
                    merged_areas.append((min_r, min_c, max_r, max_c, top_val))
            except Exception:
                pass

            def merged_value(r: int, c: int, original: str) -> str:
                for r0, c0, r1, c1, topv in merged_areas:
                    if r0 <= r <= r1 and c0 <= c <= c1:
                        return topv if topv != "" else original
                return original

            # Заголовок листа
            text_lines.append(f"## Лист: {ws.title}")

            # Таблица
            header_printed = False
            for r_idx, row in enumerate(
                ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col),
                start=min_row
            ):
                vals = []
                for c_idx, c in enumerate(row, start=min_col):
                    cell_text = esc_md_cell(fmt_cell(ws, c))
                    cell_text = merged_value(r_idx, c_idx, cell_text)
                    # гиперссылки -> markdown
                    try:
                        if getattr(c, "hyperlink", None) and cell_text:
                            href = c.hyperlink.target
                            if href:
                                cell_text = f"[{cell_text}]({href})"
                    except Exception:
                        pass
                    vals.append(cell_text)

                if not header_printed and all(v == "" for v in vals):
                    continue

                if not header_printed:
                    text_lines.append(" | ".join(vals))
                    text_lines.append(" | ".join(["---"] * len(vals)))
                    header_printed = True
                else:
                    text_lines.append(" | ".join(vals))

            # Галерея картинок для конкретного листа
            if WANT_GALLERY:
                sheet_media = sheet_to_media.get(ws.title) or []
                if sheet_media:
                    text_lines.append("")
                    text_lines.append(f"### Изображения листа: {ws.title}")
                    for i, rel in enumerate(sheet_media, 1):
                        text_lines.append(f"![{ws.title} img {i:03d}](/media/{rel})")

            text_lines.append("")

        # Фолбэк: если сопоставление не удалось, но картинки есть
        if WANT_GALLERY and not any(sheet_to_media.values()) and images:
            text_lines.append("## Изображения (из книги)")
            for i, im in enumerate(images, 1):
                text_lines.append(f"![img_{i:03d}](/media/{im['rel']})")
            text_lines.append("")

        return "\n".join(text_lines), images



    def _read_docx(self, path: pathlib.Path) -> tuple[str, list[dict]]:
        if docx is None:
            raise RuntimeError("python-docx не установлен в текущем окружении")

        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table, _Cell
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn

        d = docx.Document(str(path))
        lines: list[str] = []

        # каталог медиа для этого файла
        media_key = self._media_key_for(path)
        media_root = MEDIA_ROOT
        media_root.mkdir(parents=True, exist_ok=True)

        images: list[dict] = []
        img_idx = 0

        MIME_EXT = {
            "image/jpeg": ".jpg",
            "image/pjpeg": ".jpg",
            "image/png": ".png",
            "image/gif": ".gif",
            "image/webp": ".webp",
            "image/tiff": ".tif",
            "image/svg+xml": ".svg",
            "image/x-emf": ".emf",
            "image/emf": ".emf",
            "image/x-wmf": ".wmf",
            "image/bmp": ".bmp",
        }
        def _save_part_image(part) -> str | None:
            nonlocal img_idx
            data = getattr(part, "blob", None)
            if not data:
                return None
            mime = getattr(part, "content_type", None) or "application/octet-stream"
            ext = MIME_EXT.get(mime) or (mimetypes.guess_extension(mime) or ".bin")
            img_idx += 1
            rel = f"{media_key}/img_{img_idx:03d}{ext}"
            full = media_root / rel
            full.parent.mkdir(parents=True, exist_ok=True)
            with open(full, "wb") as f:
                f.write(data)
            sha = hashlib.sha256(data).hexdigest()
            images.append({"rel": rel.replace("\\", "/"), "mime": mime, "sha256": sha})
            return f"![img_{img_idx:03d}](/media/{rel.replace('\\', '/')})"

        def _extract_par_images(p: Paragraph) -> list[str]:
            toks = []
            # 1) Современные рисунки (DrawingML)
            try:
                for bl in p._p.xpath(".//*[local-name()='blip']"):
                    rid = bl.get(qn('r:embed')) or bl.get(qn('r:link'))
                    if not rid:
                        continue
                    part = p.part.related_parts.get(rid)
                    if part:
                        tok = _save_part_image(part)
                        if tok:
                            toks.append(tok)
            except Exception:
                pass

            # 2) Старые VML-рисунки (на всякий случай)
            try:
                for vi in p._p.xpath(".//*[local-name()='imagedata']"):
                    rid = vi.get(qn('r:id'))
                    if not rid:
                        continue
                    part = p.part.related_parts.get(rid)
                    if part:
                        tok = _save_part_image(part)
                        if tok:
                            toks.append(tok)
            except Exception:
                pass

            return toks

        # --- остальной код из вашей версии (_heading_level, list_info, render_runs, iter_block_items) ---
        def heading_level(p: Paragraph) -> int:
            name = (getattr(p.style, "name", "") or "").lower()
            m = re.search(r'(heading|заголовок)\s*(\d+)', name)
            if m:
                return min(6, int(m.group(2)))
            try:
                el = p._element
                ol = el.xpath(".//w:outlineLvl", namespaces=el.nsmap)
                if ol:
                    val = getattr(ol[0], "val", None)
                    if val is not None:
                        return min(6, int(val) + 1)
            except Exception:
                pass
            return 0

        def list_info(p: Paragraph):
            try:
                numPr = p._element.pPr.numPr
            except Exception:
                numPr = None
            if numPr is None:
                return None
            ilvl = getattr(numPr.ilvl, "val", 0)
            numId = getattr(numPr.numId, "val", None)
            try:
                ilvl = int(ilvl) if ilvl is not None else 0
            except Exception:
                ilvl = 0
            try:
                numId = int(numId) if numId is not None else None
            except Exception:
                numId = None
            nm = (getattr(p.style, "name", "") or "").lower()
            is_ordered = ("number" in nm) or ("нумер" in nm)
            return ilvl, numId, is_ordered

        def render_runs(p: Paragraph) -> str:
            out = []
            for r in p.runs:
                t = (r.text or "")
                if not t:
                    continue
                sname = (getattr(r.style, "name", "") or "").lower()
                if getattr(r, "bold", False):
                    t = f"**{t}**"
                if getattr(r, "italic", False):
                    t = f"*{t}*"
                if "code" in sname or "кода" in sname:
                    t = f"`{t}`"
                out.append(t)
            text_joined = "".join(out).strip()

            # В конце параграфа добавляем маркеры для картинок, если они были
            img_toks = _extract_par_images(p)
            if img_toks:
                # Если в параграфе был текст — перенос и картинки, иначе только картинки
                if text_joined:
                    text_joined = text_joined + ("\n" + "\n".join(img_toks))
                else:
                    text_joined = "\n".join(img_toks)

            return text_joined if text_joined else (p.text or "").strip()

        def esc_md_cell(s: str) -> str:
            return (s or "").replace("|", "\\|").strip()

        counters: dict[tuple[int | None, int], int] = {}
        def next_number(num_id: int | None, lvl: int) -> int:
            key = (num_id, lvl)
            counters[key] = counters.get(key, 0) + 1
            return counters[key]

        def iter_block_items(parent):
            if isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                parent_elm = parent.element.body
            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        for block in iter_block_items(d):
            if isinstance(block, Paragraph):
                t = render_runs(block)
                if not t:
                    continue
                h = heading_level(block)
                if h > 0:
                    lines.append(f'{"#" * h} {t}')
                    continue
                li = list_info(block)
                if li is not None:
                    ilvl, numId, is_ord = li
                    indent = "    " * max(ilvl, 0)
                    if is_ord:
                        n = next_number(numId, ilvl)
                        lines.append(f"{indent}{n}. {t}")
                    else:
                        lines.append(f"{indent}- {t}")
                else:
                    lines.append(t)

            elif isinstance(block, Table):
                if lines and lines[-1] != "":
                    lines.append("")
                rows = [[esc_md_cell(c.text) for c in row.cells] for row in block.rows]
                rows = [r for r in rows if any(cell for cell in r)]
                if not rows:
                    continue
                header = rows[0]
                lines.append(" | ".join(header))
                lines.append(" | ".join(["---"] * len(header)))
                for r in rows[1:]:
                    lines.append(" | ".join(r))
                lines.append("")

        return "\n".join(lines), images


    
    BOILERPLATE_CANDIDATES = {
        "в отчёте видно:",
        "пример строки материал:",
    }

    def _strip_wrapping_quotes(self, text: str) -> str:
        QUOTE_PAIRS = [('"', '"'), ('«','»'), ('“','”'), ('‘','’')]
        lines = []
        for ln in text.splitlines():
            s = ln.strip()
            for ql, qr in QUOTE_PAIRS:
                if len(s) >= 2 and s.startswith(ql) and s.endswith(qr):
                    s = s[1:-1].strip()
                    break
            lines.append(s)
        return "\n".join(lines)

    def _collapse_consecutive_duplicates(self, text: str) -> str:
        # схлопываем подряд идущие одинаковые строки
        out = []
        prev_key = None
        for ln in text.splitlines():
            key = self._normalize_like_db(ln)
            if key == prev_key:
                continue
            out.append(ln)
            prev_key = key
        return "\n".join(out)

    def _drop_high_freq_boilerplate(self, text: str, min_repeat: int = 2) -> str:
        paras = re.split(r"\n\s*\n+", text)
        norm = lambda s: self._normalize_like_db(s)
        counts = {}
        for p in paras:
            if not p.strip():    # ← пропустить пустые
                continue
            k = norm(p)
            counts[k] = counts.get(k, 0) + 1

        out, seen = [], set()
        for p in paras:
            if not p.strip():
                out.append(p)
                continue
            k = norm(p)
            is_short = len(p) < 200
            looks_boiler = any(norm(bp) in k for bp in self.BOILERPLATE_CANDIDATES)
            if (counts.get(k, 0) >= min_repeat and is_short) or looks_boiler:
                if k in seen:
                    continue
                seen.add(k)
            out.append(p)
        return "\n\n".join(out)
    
    def clean_text(self, text: str) -> str:
        """Улучшенная очистка текста"""
        # Нормализация пробелов и символов
        text = re.sub(r'[\xa0\u200b\u00a0\ufeff]', ' ', text)
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Сохраняем структурные элементы
        text = re.sub(r'\n{4,}', '\n\n\n', text)  # Максимум 3 переноса
        text = re.sub(r'[ \t]+', ' ', text)  # Нормализация пробелов
        text = re.sub(r'[ \t]+\n', '\n', text)  # Удаление пробелов в конце строк
        
        # Удаление артефактов пагинации
        text = re.sub(r'Стр(?:аница)?\.?\s*\d+\s*(?:из\s*\d+)?', '', text, flags=re.IGNORECASE)
        text = re.sub(r'_{10,}', '', text)  # Длинные подчеркивания
        text = re.sub(r'-{10,}', '', text)  # Длинные тире
        
        # Исправление типографики
        text = re.sub(r'([а-яё])\s*-\s*\n\s*([а-яё])', r'\1\2', text, flags=re.IGNORECASE)  # Переносы слов
        text = re.sub(r'(\w)\s*\n\s*(\w)', r'\1 \2', text)  # Склейка разорванных предложений
        
        # 1) убрать кавычки у строк
        text = self._strip_wrapping_quotes(text)
        # 2) схлопнуть подряд идущие дубликаты строк
        text = self._collapse_consecutive_duplicates(text)
        # 3) убрать высокочастотный boilerplate и повторы коротких абзацев
        text = self._drop_high_freq_boilerplate(text, min_repeat=2)

        return text.strip()
    
    def semantic_chunk(self, text: str) -> List[Tuple[str, Dict]]:
        """Семантическое разбиение с учетом структуры"""
        chunks_with_meta: List[Tuple[str, Dict]] = []

        # Разбиваем на секции по заголовкам
        sections = self._split_by_headers(text)

        for section_title, section_text, level in sections:
            section_text = section_text.strip()

            # 1) Отсечь оглавление/пустышки: если текст совпадает с заголовком
            #    или слишком короткий — попробуем слить с предыдущим, иначе пропустим.
            same_as_title = (
                self._normalize_like_db(section_text) ==
                self._normalize_like_db(section_title)
            )
            if same_as_title or len(section_text) < 80:
                if chunks_with_meta:
                    prev_chunk, prev_meta = chunks_with_meta[-1]
                    combined = f"{prev_chunk}\n\n{section_title}\n{section_text}".strip()
                    if self._tok_len(combined) <= self.max_chunk_tokens:
                        chunks_with_meta[-1] = (combined, prev_meta)
                # если слить не получилось (или это первая секция) — просто пропустим
                continue

            # 2) Оценка важности после фильтра пустышек
            importance = self._calculate_importance(section_text)

            # 3) Длинные секции — дробим
            sec_toks = self._tok_len(section_text)
            if sec_toks > self.max_chunk_tokens:
                sub_chunks = self._split_long_section(section_text)
                for i, chunk in enumerate(sub_chunks):
                    meta = {
                        'section_title': section_title,
                        'section_level': level,
                        'chunk_index': i,
                        'total_chunks': len(sub_chunks),
                        'importance': importance
                    }
                    chunks_with_meta.append((chunk, meta))

            # 4) Нормальные секции — кладем как есть
            elif sec_toks >= self.min_chunk_tokens:
                meta = {
                    'section_title': section_title,
                    'section_level': level,
                    'chunk_index': 0,
                    'total_chunks': 1,
                    'importance': importance
                }
                chunks_with_meta.append((section_text, meta))

            # 5) Небольшие — пробуем присоединить к предыдущему
            elif chunks_with_meta and sec_toks > max(30, self.min_chunk_tokens // 3):
                prev_chunk, prev_meta = chunks_with_meta[-1]
                combined = f"{prev_chunk}\n\n{section_title}\n{section_text}"
                if self._tok_len(combined) <= self.max_chunk_tokens:
                    chunks_with_meta[-1] = (combined, prev_meta)
                else:
                    meta = {
                        'section_title': section_title,
                        'section_level': level,
                        'chunk_index': 0,
                        'total_chunks': 1,
                        'importance': importance
                    }
                    chunks_with_meta.append((section_text, meta))
            # иначе — пропускаем мелочь

        return chunks_with_meta
    
    def _split_by_headers(self, text: str) -> List[Tuple[str, str, int]]:
        """Разбиение текста по заголовкам с определением уровня"""
        sections = []
        lines = text.split('\n')
        current_section = []
        current_title = "Введение"
        current_level = 0
        
        for line in lines:
            toc_like = re.match(r'^\d+(?:\.\d+)*\s+.+\s\d{1,4}$', line.strip())
            if toc_like:
                # это строка оглавления вида "8. Название ... 4" — не считаем заголовком
                current_section.append(line)
                continue

            raw = line.strip()
            # Не считать длинные "1. …" списочные фразы заголовками
            if re.match(r'^\d+(?:\.\d+)*\s+[А-ЯA-Z]', raw) and len(raw) > 100:
                current_section.append(line)
                continue

            # Проверяем, является ли строка заголовком
            is_header = False
            for pattern, level in self.header_patterns:
                if re.match(pattern, line.strip()) \
                and not re.match(r'^\d+(?:\.\d+)*\s+.+\s\d{1,4}$', line.strip()):
                    # Сохраняем предыдущую секцию
                    if current_section:
                        section_text = '\n'.join(current_section).strip()
                        if section_text:
                            sections.append((current_title, section_text, current_level))
                    
                    # Начинаем новую секцию
                    current_title = line.strip()
                    current_level = level
                    current_section = []
                    is_header = True
                    break
            
            if not is_header:
                current_section.append(line)
        
        # Добавляем последнюю секцию
        if current_section:
            section_text = '\n'.join(current_section).strip()
            if section_text:
                sections.append((current_title, section_text, current_level))
        
        return sections
    
    def _split_long_section(self, text: str) -> List[str]:
        """Разбивка длинной секции по ТОКЕНАМ: параграфы → предложения при переполнении → overlap."""
        if not getattr(self, "tokenizer", None):
            return self._split_long_section_charwise(text)

        max_toks = self.chunk_tokens
        overlap_toks = self.overlap_tokens
        hard_char_cap = getattr(self, "hard_char_cap", 4000)

        chunks: List[str] = []
        buf: List[str] = []
        buf_toks = 0

        paragraphs = re.split(r"\n{2,}", (text or "").strip())
        for para in paragraphs:
            if not para.strip():
                continue

            pt = self._tok_len(para)

            # Слишком длинный параграф — дробим по предложениям
            if pt > max_toks:
                if buf:
                    chunks.append("\n\n".join(buf))
                    buf, buf_toks = [], 0
                sents = self._split_into_sentences(para)
                chunks.extend(self._split_sentences_with_overlap(sents, max_toks, overlap_toks))
                continue

            # Жадно пакуем параграфы
            if buf_toks + pt <= max_toks:
                buf.append(para)
                buf_toks += pt
            else:
                if buf:
                    chunks.append("\n\n".join(buf))
                # хвост, который и overlap соблюдёт (по возможности), и с новым параграфом влезет в max_toks
                tail = self._tail_fit_for_overlap(buf, overlap_toks, next_len=pt, max_toks=max_toks) if overlap_toks > 0 else []
                buf = tail + [para]
                buf_toks = sum(self._tok_len(x) for x in buf)

        if buf:
            chunks.append("\n\n".join(buf))

        # Жёсткая отсечка по символам (редкие случаи очень больших чанков)
        safe_chunks: List[str] = []
        for ch in chunks:
            if len(ch) <= hard_char_cap:
                safe_chunks.append(ch)
                continue

            # режем по пробельным символам, не превышая hard_char_cap
            start = 0
            n = len(ch)
            while start < n:
                end = min(n, start + hard_char_cap)
                cut = end
                if end < n:
                    window = ch[start:end]
                    # найдём ПОСЛЕДНИЙ пробельный символ в окне
                    last_ws = None
                    for m in re.finditer(r'\s', window):
                        last_ws = m
                    if last_ws is not None:
                        cut = start + last_ws.start()

                piece = ch[start:cut].strip()
                if piece:
                    safe_chunks.append(piece)
                # гарантия прогресса даже если пробела не нашли
                start = cut if cut > start else end

        return safe_chunks


    def _tail_fit_for_overlap(self, parts: List[str], overlap_toks: int, next_len: int, max_toks: int) -> List[str]:
        """
        Берём хвост из предыдущего буфера так, чтобы:
        1) по возможности набрать >= overlap_toks,
        2) суммарно (tail + next_len) <= max_toks.
        Если требования конфликтуют, приоритет у (2), overlap уменьшается.
        """
        if not parts:
            return []
        tail: List[str] = []
        tail_len = 0
        # идём с конца вперёд и добавляем, пока влезает
        for p in reversed(parts):
            l = self._tok_len(p)
            # если добавление p нарушит бюджет с новым параграфом — стоп
            if tail_len + l + next_len > max_toks:
                break
            tail.insert(0, p)
            tail_len += l
            # если перекрытие уже достаточно — можно остановиться
            if tail_len >= overlap_toks:
                break
        return tail

    def _split_sentences_with_overlap(self, sentences: List[str], max_toks: int, overlap_toks: int) -> List[str]:
        """
        Разбиваем список предложений на чанки с перекрытием по токенам.
        Сверхдлинные предложения режем по пробельным токенам.
        """
        def _yield_long_sentence_parts(s: str) -> List[str]:
            toks_limit = max_toks
            parts: List[str] = []
            cur, cur_len = [], 0
            # делим по «словам/пробелам», чтобы не резать середину слов
            for piece in re.split(r'(\s+)', s):
                if not piece:
                    continue
                l = self._tok_len(piece)
                if cur_len + l > toks_limit and cur:
                    parts.append("".join(cur).strip())
                    cur, cur_len = [piece], l
                else:
                    cur.append(piece)
                    cur_len += l
            if cur:
                parts.append("".join(cur).strip())
            return [p for p in parts if p]

        chunks: List[str] = []
        cur: List[str] = []
        cur_len = 0

        i = 0
        while i < len(sentences):
            s = sentences[i]
            s_len = self._tok_len(s)

            # если предложение само слишком длинное — нарезаем на куски и обрабатываем их последовательно
            if s_len > max_toks:
                parts = _yield_long_sentence_parts(s)
                # подставим куски вместо исходного предложения
                sentences = sentences[:i] + parts + sentences[i+1:]
                # не увеличиваем i — на следующей итерации возьмём первый part
                continue

            if cur_len + s_len <= max_toks:
                cur.append(s)
                cur_len += s_len
                i += 1
                continue

            # нужно закрывать текущий чанк и начинать новый с перекрытием
            if cur:
                chunks.append(" ".join(cur))

            # соберём хвост перекрытия, чтобы новый чанк вместе с s влез в max_toks
            tail: List[str] = []
            tail_len = 0
            for t in reversed(cur):
                l = self._tok_len(t)
                if tail_len + l + s_len > max_toks:
                    break
                tail.insert(0, t)
                tail_len += l
                if tail_len >= overlap_toks:
                    break

            cur = tail + [s]
            cur_len = sum(self._tok_len(x) for x in cur)
            i += 1

        if cur:
            chunks.append(" ".join(cur))
        return chunks

    def _split_long_section_charwise(self, text: str) -> List[str]:
        """Разбиение длинной секции с сохранением семантической целостности"""
        chunks = []
        
        # Сначала пробуем разбить по абзацам
        paragraphs = re.split(r'\n\n+', text)
        current_chunk = []
        current_size = 0
        
        for para in paragraphs:
            para_size = len(para)
            
            # Если абзац сам по себе слишком большой
            if para_size > self.chunk_size:
                # Сохраняем текущий чанк
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = []
                    current_size = 0
                
                # Разбиваем большой абзац по предложениям
                sentences = self._split_into_sentences(para)
                for sent_chunk in self._group_sentences(sentences):
                    chunks.append(sent_chunk)
            
            # Если добавление абзаца превысит лимит
            elif current_size + para_size > self.chunk_size:
                # Сохраняем текущий чанк
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                current_chunk = [para]
                current_size = para_size
            
            # Добавляем абзац к текущему чанку
            else:
                current_chunk.append(para)
                current_size += para_size
        
        # Добавляем последний чанк
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Умное разбиение на предложения"""
        # Защищаем сокращения
        text = re.sub(r'\b([А-Я])\.\s*([А-Я])\.', r'\1·\2·', text)  # И.О. Фамилия
        text = re.sub(r'\b(т|ст|п|пп|гл|разд|прил|рис|табл|см|ср|напр|др|пр)\.\s*', r'\1·', text, flags=re.IGNORECASE)
        
        # Разбиваем по знакам препинания
        sentences = re.split(r'([.!?])\s+', text)
        
        # Собираем предложения обратно
        result = []
        current = ""
        for i, part in enumerate(sentences):
            if part in '.!?':
                current += part
                if i + 1 < len(sentences):
                    # Проверяем, не начинается ли следующая часть с маленькой буквы
                    if not re.match(r'^[а-яё]', sentences[i + 1]):
                        result.append(current.strip())
                        current = ""
            else:
                current += " " + part if current else part
        
        if current:
            result.append(current.strip())
        
        # Восстанавливаем точки
        result = [s.replace('·', '.') for s in result if s]
        
        return result
    
    def _group_sentences(self, sentences: List[str]) -> List[str]:
        """Группировка предложений в чанки оптимального размера"""
        chunks = []
        current = []
        current_size = 0
        
        for sent in sentences:
            sent_size = len(sent)
            
            if current_size + sent_size > self.chunk_size and current:
                chunks.append(' '.join(current))
                # Добавляем overlap - последнее предложение предыдущего чанка
                if self.overlap > 0 and len(current) > 1:
                    current = [current[-1], sent]
                    current_size = len(current[-1]) + sent_size
                else:
                    current = [sent]
                    current_size = sent_size
            else:
                current.append(sent)
                current_size += sent_size
        
        if current:
            chunks.append(' '.join(current))
        
        return chunks
    
    def _take_tail_for_overlap(self, parts: List[str], overlap_toks: int) -> List[str]:
        """Берём хвост предыдущего чанка по токенам для перекрытия."""
        tail: List[str] = []
        acc = 0
        for p in reversed(parts):
            tail.insert(0, p)
            acc += self._tok_len(p)
            if acc >= overlap_toks:
                break
        return tail

    
    def _calculate_importance(self, text: str) -> float:
        """Расчет важности текста"""
        text_lower = text.lower()
        score = 1.0
        
        # Проверяем наличие важных терминов
        for term, weight in self.important_terms.items():
            if term in text_lower:
                score += weight * 0.1
        
        # Учитываем наличие списков
        if re.search(r'^\s*[-•*]\s+', text, re.MULTILINE):
            score += 0.2
        
        # Учитываем наличие примеров кода
        if re.search(r'```|^\s{4,}\S', text, re.MULTILINE):
            score += 0.3
        
        # Учитываем наличие таблиц
        if '|' in text and text.count('|') > 5:
            score += 0.3
        
        return min(score, 2.0)  # Максимальная важность 2.0
    
    def extract_advanced_metadata(self, text: str, filepath: pathlib.Path) -> Dict:
        """Расширенное извлечение метаданных"""
        meta = {
            'source_file': filepath.name,
            'file_size': filepath.stat().st_size,
            'encoding': self._detect_encoding(filepath)
        }
        
        # Извлечение версии
        version_patterns = [
            r'[Вв]ерсия\s*([\d.]+)',
            r'v([\d.]+)',
            r'Version\s*([\d.]+)',
            r'Релиз\s*([\d.]+)'
        ]
        for pattern in version_patterns:
            match = re.search(pattern, text[:1000], re.IGNORECASE)
            if match:
                meta['version'] = match.group(1)
                break
        
        # Извлечение дат
        date_patterns = [
            (r'(\d{2})\.(\d{2})\.(\d{4})', 'dd.mm.yyyy'),
            (r'(\d{4})-(\d{2})-(\d{2})', 'yyyy-mm-dd'),
            (r'(\d{2})/(\d{2})/(\d{4})', 'dd/mm/yyyy')
        ]
        for pattern, format_type in date_patterns:
            match = re.search(pattern, text[:1000])
            if match:
                meta['date'] = match.group(0)
                meta['date_format'] = format_type
                break
        
        # Извлечение ключевых слов через частотный анализ
        words = re.findall(r'\b[а-яёА-ЯЁa-zA-Z]{4,}\b', text)
        word_freq = Counter(word.lower() for word in words)
        
        # Фильтруем стоп-слова
        stop_words = {'этот', 'который', 'такой', 'только', 'также', 'более', 'менее', 'через', 'после', 'перед'}
        keywords = [word for word, freq in word_freq.most_common(20) 
                   if word not in stop_words and freq > 3]
        
        if keywords:
            meta['keywords'] = keywords[:10]
        
        # Определение типа документа
        doc_type_indicators = {
            'руководство': ['настройка', 'установка', 'конфигурация'],
            'справочник': ['функция', 'метод', 'параметр', 'свойство'],
            'инструкция': ['шаг', 'действие', 'выполните', 'нажмите'],
            'описание': ['модуль', 'компонент', 'архитектура', 'структура']
        }
        
        text_lower = text[:3000].lower()
        for doc_type, indicators in doc_type_indicators.items():
            if sum(1 for ind in indicators if ind in text_lower) >= 2:
                meta['doc_type'] = doc_type
                break
        
        # Статистика документа
        meta['stats'] = {
            'total_chars': len(text),
            'total_words': len(words),
            'total_lines': text.count('\n') + 1,
            'avg_word_length': sum(len(w) for w in words) / len(words) if words else 0
        }
        
        return meta
    
    def _detect_encoding(self, filepath: pathlib.Path) -> str:
        """Определение кодировки файла"""
        content = filepath.read_bytes()[:1000]
        
        if content.startswith(b'\xef\xbb\xbf'):
            return 'utf-8-sig'
        elif content.startswith(b'\xff\xfe'):
            return 'utf-16-le'
        elif content.startswith(b'\xfe\xff'):
            return 'utf-16-be'
        
        # Эвристическое определение
        try:
            content.decode('utf-8')
            return 'utf-8'
        except:
            try:
                content.decode('cp1251')
                return 'cp1251'
            except:
                return 'unknown'
    
    def process_txt(self, filepath: pathlib.Path) -> List[Document]:
        logger.info(f"Обработка файла: {filepath}")

        suffix = filepath.suffix.lower()
        if suffix == ".docx":
            try:
                raw_text, images = self._read_docx(filepath)
            except Exception as e:
                logger.error(f"Ошибка чтения DOCX {filepath}: {e}")
                return []
        elif suffix == ".xlsx":
            try:
                raw_text, images = self._read_xlsx(filepath)
            except Exception as e:
                logger.error(f"Ошибка чтения XLSX {filepath}: {e}")
                return []
        elif suffix == ".csv":
            try:
                raw_text, images = self._read_csv(filepath)
            except Exception as e:
                logger.error(f"Ошибка чтения CSV {filepath}: {e}")
                return []
        elif suffix == ".pdf":
            try:
                raw_text, images = self._read_pdf(filepath)
            except Exception as e:
                logger.error(f"Ошибка чтения PDF {filepath}: {e}")
                return []
        else:
            raw_text = self.read_text_with_encoding(filepath)
            images = []

        if suffix in (".txt", ".md"):
            text = self.clean_text(raw_text)
        else:
            text = raw_text
            
        base_meta = self.extract_advanced_metadata(text, filepath)

        # PDF-метаданные (если есть) добавим в meta
        if suffix == ".pdf":
            try:
                pdf_meta = self._extract_pdf_metadata(filepath)
                if pdf_meta:
                    base_meta.update(pdf_meta)
            except Exception:
                pass

        if images:
            base_meta["images"] = images

        chunks_with_meta = self.semantic_chunk(text)
        
        documents = []
        groups: Dict[str, List[Document]] = {}

        for i, (chunk_text, chunk_meta) in enumerate(chunks_with_meta, 1):
            section_title = chunk_meta.get('section_title', 'Общие сведения')
            importance = chunk_meta.get('importance', 1.0)

            doc = Document(
                book=filepath.stem,
                section=section_title,
                page=i,
                text=chunk_text,
                meta={**base_meta, **chunk_meta},
                doc_hash=self.compute_hash(chunk_text),
                importance_score=importance
            )
            documents.append(doc)
            groups.setdefault(section_title, []).append(doc)

        # вставляем родителя и проставляем parent_* детям
        parent_docs: List[Document] = []
        for title, childs in groups.items():
            pg = AdvancedDocumentProcessor._parent_group_key(filepath.stem, title)
            # короткий «аннотационный» текст родителя: заголовок + первые 800 символов секции
            parent_text = f"{title}\n\n" + ' '.join(c.text for c in childs)[:800]

            parent = Document(
                book=filepath.stem,
                section=title,
                page=None,  # NULL
                text=parent_text,
                meta={**base_meta, 'section_title': title, 'section_level': childs[0].meta.get('section_level', 0)},
                doc_hash=self.compute_hash(parent_text),
                importance_score=max(c.importance_score for c in childs),
                is_parent=True,
                parent_group=pg,
                child_index=None,
                parent_title=title
            )
            if childs and isinstance(childs[0].page, int):
                parent.page = childs[0].page

            parent.meta['is_parent'] = True    

            parent_docs.append(parent)

            # детям — ссылка на группу и порядковый индекс
            for idx, c in enumerate(childs):
                c.is_parent = False
                c.parent_group = pg
                c.child_index = idx
                c.parent_title = title
                c.meta.update({"parent_group": pg, "child_index": idx, "parent_title": title})

        # Родители вперёд не обязательно, но удобно
        out = parent_docs + documents
        logger.info(f"Извлечено {len(out)} фрагментов (вместе с родителями) из {filepath.name}")
        return out
    
    @staticmethod
    def _unaccent_basic(s: str) -> str:
        s = unicodedata.normalize('NFKD', s)
        return ''.join(ch for ch in s if unicodedata.category(ch) != 'Mn')
    
    @staticmethod
    def _normalize_like_db(s: str) -> str:
        s = (s or '').lower().replace('ё','е')
        s = AdvancedDocumentProcessor._unaccent_basic(s)
        s = s.translate({cp: ' ' for cp in SPACE_MAP})
        s = ''.join(ch for ch in s if ord(ch) not in INVISIBLES)
        return re.sub(r'\s+', ' ', s).strip()
    
    @staticmethod
    def _parent_group_key(book: str, title: str) -> str:
        # стабильный ключ группы: md5 от нормализованного "book::title"
        base = f"{AdvancedDocumentProcessor._normalize_like_db(book)}::{AdvancedDocumentProcessor._normalize_like_db(title)}"
        return hashlib.md5(base.encode('utf-8')).hexdigest()
    
    @staticmethod
    def compute_hash(text: str) -> str:
        normalized = AdvancedDocumentProcessor._normalize_like_db(text)
        return hashlib.md5(normalized.encode('utf-8')).hexdigest()
    
    def read_text_with_encoding(self, filepath: pathlib.Path) -> str:
        """Чтение файла с автоопределением кодировки"""
        content = filepath.read_bytes()
        
        # Проверка на BOM
        if content.startswith(b'\xef\xbb\xbf'):
            return content[3:].decode('utf-8')
        elif content.startswith(b'\xff\xfe'):
            return content[2:].decode('utf-16-le')
        elif content.startswith(b'\xfe\xff'):
            return content[2:].decode('utf-16-be')
        
        # Порядок кодировок для Windows
        encodings = [
            'utf-8', 'cp1251', 'cp866', 'windows-1251',
            'utf-8-sig', 'utf-16', 'koi8-r', 'iso-8859-5'
        ]
        
        for encoding in encodings:
            try:
                decoded = content.decode(encoding)
                if decoded and not any(ord(c) > 0xFFFF for c in decoded):
                    logger.debug(f"Файл {filepath.name} декодирован с {encoding}")
                    return decoded
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        logger.warning(f"Не удалось определить кодировку для {filepath}")
        return content.decode('utf-8', errors='ignore')
    
    def to_vec_literal(self, vec) -> str:
        """Преобразование вектора в строку для PostgreSQL"""
        return "[" + ",".join(f"{float(x):.8f}" for x in vec) + "]"
    
    def save_to_db(self, documents: List[Document], batch_size: int = 16, clean_start: bool = False):
        """Сохранение документов в БД, опционально с чистым стартом (TRUNCATE)."""
        logger.info(f"Сохранение {len(documents)} документов в БД... (clean_start={clean_start})")
        if not documents:
            logger.warning("Нет документов для сохранения")
            return

        # Дедупликация в текущем запуске
        unique_docs = {}
        for doc in documents:
            key = (doc.parent_group or '', bool(doc.is_parent), doc.doc_hash)
            if key not in unique_docs:
                unique_docs[key] = doc
        documents = list(unique_docs.values())
        logger.info(f"После дедупликации: {len(documents)} уникальных документов")

        # --- Префильтр по тройке (parent_group, is_parent, doc_hash) в БД ---
        from io import StringIO
        buf = StringIO(''.join(
            f"{(d.parent_group or '')}\t{str(bool(d.is_parent)).lower()}\t{d.doc_hash}\n"
            for d in documents
        ))

        with psycopg2.connect(**self.db_config) as pre_conn, pre_conn.cursor() as pre_cur:
            pre_conn.set_client_encoding('UTF8')
            pre_cur.execute("CREATE TEMP TABLE keys(pg text, ip boolean, h text) ON COMMIT DROP;")
            pre_cur.copy_expert("COPY keys FROM STDIN WITH (FORMAT text)", buf)
            pre_cur.execute("""
                SELECT k.pg, k.ip, k.h
                FROM keys k
                JOIN public.docs d
                ON COALESCE(d.parent_group,'') = k.pg
                AND d.is_parent = k.ip
                AND d.doc_hash  = k.h
            """)
            existing = {(pg, ip, h) for pg, ip, h in pre_cur.fetchall()}

        before = len(documents)
        documents = [d for d in documents
                    if ((d.parent_group or ''), bool(d.is_parent), d.doc_hash) not in existing]
        logger.info("Отфильтровано по БД ДО эмбеддинга: %d; к эмбеддингу пойдёт %d",
                    before - len(documents), len(documents))
        if not documents:
            logger.info("Все документы уже в БД — ничего эмбеддить/вставлять")
            return
        # --- /префильтр ---




        # Генерация эмбеддингов
        texts = [doc.text for doc in documents]
        embeddings = []
        dim: Optional[int] = None
        for i in tqdm(range(0, len(texts), batch_size), desc="Генерация эмбеддингов"):
            batch = [f"passage: {t}" for t in texts[i:i + batch_size]]
            try:
                batch_emb = self.model.encode(batch, batch_size=batch_size, normalize_embeddings=True, convert_to_numpy=True, show_progress_bar=False)
                for emb in batch_emb:
                    # дополнительная нормализация (на случай если модель не вернула нормированные)
                    norm = np.linalg.norm(emb)
                    embeddings.append(emb / norm if norm > 0 else emb)
                if dim is None and len(batch_emb):
                    dim = int(np.asarray(batch_emb[0]).shape[0])
            except Exception as e:
                logger.error(f"Ошибка при генерации эмбеддингов: {e}")
                for pref in batch: 
                    try:
                        emb = self.model.encode(
                            [pref],
                            batch_size=1,
                            normalize_embeddings=True,
                            convert_to_numpy=True,
                            show_progress_bar=False
                        )[0]
                        # дополнительная нормализация на всякий случай
                        norm = np.linalg.norm(emb)
                        emb = emb / norm if norm > 0 else emb
                    except Exception as ee:
                        logger.error(f"  └─ Не удалось сэмбеддить элемент: {ee}")
                        emb = np.zeros(dim or self.emb_dim, dtype=np.float32)
                    embeddings.append(emb)
                    if dim is None:
                        dim = int(np.asarray(emb).shape[0])

        # --- SANITY CHECK ПЕРЕД ЗАПИСЬЮ В БД ---
        # 0) гарантируем известную размерность
        dim = dim or self.emb_dim
        if dim is None:
            dim = int(np.asarray(embeddings[0]).shape[0]) if embeddings else 1024
            logger.warning(f"Размерность эмбеддинга не определена явно. Использую dim={dim}")

        # 1) выровнять количество эмбеддингов и документов
        if len(embeddings) != len(documents):
            logger.error(f"Мисматч: embeddings={len(embeddings)} vs documents={len(documents)}")
            if len(embeddings) < len(documents):
                pad = len(documents) - len(embeddings)
                logger.warning(f"Добиваю {pad} эмбеддингов нулями до длины документов")
                embeddings.extend([np.zeros(dim, dtype=np.float32) for _ in range(pad)])  # отдельные массивы
            else:
                logger.warning(f"Обрезаю лишние эмбеддинги: {len(embeddings)} -> {len(documents)}")
                embeddings = embeddings[:len(documents)]

        # 2) приведение формы/типа + очистка NaN/Inf + повторная нормализация
        fixed_embeddings = []
        for i, emb in enumerate(embeddings):
            e = np.asarray(emb, dtype=np.float32).reshape(-1)

            # неверная размерность → подрезать/дополнить нулями
            if e.size != dim:
                logger.warning(f"[emb#{i}] Размерность {e.size} != {dim}. Исправляю.")
                if e.size == 0:
                    e = np.zeros(dim, dtype=np.float32)
                elif e.size > dim:
                    e = e[:dim]
                else:
                    e = np.pad(e, (0, dim - e.size))

            # NaN/Inf → в нули
            if not np.all(np.isfinite(e)):
                logger.warning(f"[emb#{i}] Обнаружены NaN/Inf. Заменяю на 0.")
                e = np.nan_to_num(e, nan=0.0, posinf=0.0, neginf=0.0)

            # нормализация (если не нулевой вектор)
            n = float(np.linalg.norm(e))
            if n > 0 and not np.isclose(n, 1.0, rtol=1e-3, atol=1e-3):
                e = e / n

            fixed_embeddings.append(e)

        embeddings = fixed_embeddings
        # --- /SANITY CHECK ---

        # Подключение к БД
        conn = None
        cur = None
        try:
            conn = psycopg2.connect(**self.db_config)
            conn.set_client_encoding('UTF8')

            # Чистый старт: TRUNCATE перед вставкой
            if clean_start:
                conn.autocommit = True
                with conn.cursor() as c:
                    c.execute("TRUNCATE TABLE public.docs RESTART IDENTITY;")
                logger.info("✓ Таблица docs очищена (TRUNCATE + RESTART IDENTITY)")
                conn.autocommit = False

            cur = conn.cursor()

            # Настройки для массовой вставки в рамках этой транзакции
            cur.execute("SET LOCAL synchronous_commit = OFF;")
            cur.execute("SET LOCAL work_mem = '256MB';")

            zero_count = sum(1 for e in embeddings if not np.any(e))
            if zero_count:
                logger.warning(f"{zero_count} нулевых эмбеддингов (из {len(embeddings)}).")
 
            # Подготовка данных
            rows = []
            for doc, emb in zip(documents, embeddings):
                meta_json = json.dumps(doc.meta or {}, ensure_ascii=False)
                vec_literal = self.to_vec_literal(emb)
                rows.append((
                    doc.book, doc.section, doc.page, doc.text, doc.doc_hash, vec_literal, meta_json,
                    doc.importance_score,
                    't' if doc.is_parent else 'f',
                    doc.parent_group,
                    doc.child_index,
                    doc.parent_title
                ))

            if not rows:
                logger.info("Нечего вставлять (всё уже есть)")
                conn.commit()
            else:
                from io import StringIO
                buffer = StringIO()


                def _copy_escape(val):
                    if val is None:
                        return r'\N'
                    if isinstance(val, str):
                        # важный шаг — убрать NUL
                        val = val.replace('\x00', '')
                        # стандартные экранирования для COPY text
                        val = (val.replace('\\', '\\\\')
                                .replace('\t', '\\t')
                                .replace('\n', '\\n')
                                .replace('\r', '\\r'))
                        return val
                    return str(val)

                for row in rows:
                    buffer.write('\t'.join(_copy_escape(v) for v in row) + '\n')
                buffer.seek(0)
                
                cur.execute(f"""
                    CREATE TEMP TABLE docs_stage(
                        book text, section text, page integer, text text,
                        doc_hash text,
                        emb vector({dim}), meta jsonb, importance_score real,
                        is_parent boolean, parent_group text, child_index integer, parent_title text
                    ) ON COMMIT DROP;
                """)

                cur.copy_expert("""
                    COPY docs_stage(
                        book, section, page, text, doc_hash,
                        emb, meta, importance_score,
                        is_parent, parent_group, child_index, parent_title
                    )
                    FROM STDIN WITH (FORMAT text, NULL '\\N')
                """, buffer)

                cur.execute("""
                    INSERT INTO public.docs(
                        book, section, page, text, doc_hash, emb, meta, importance_score,
                        is_parent, parent_group, child_index, parent_title
                    )
                    SELECT
                        s.book, s.section, s.page, s.text, s.doc_hash, s.emb, s.meta, s.importance_score,
                        s.is_parent, COALESCE(s.parent_group, ''), s.child_index, s.parent_title
                    FROM docs_stage s
                    ON CONFLICT (parent_group, is_parent, doc_hash) DO NOTHING
                    RETURNING id;
                """)
                inserted = cur.rowcount  # в psycopg2 для вставки с RETURNING это корректно
                conn.commit()
                logger.info("✓ Добавлено %d новых документов", inserted)


            # Анализ — вне транзакции
            conn.autocommit = True
            with conn.cursor() as c:
                c.execute("VACUUM ANALYZE public.docs;")
            logger.info("✓ VACUUM ANALYZE public.docs выполнен")

        except Exception as e:
            logger.error(f"Ошибка при работе с БД: {e}")
            if conn:
                conn.rollback()
            raise
        finally:
            if cur:
                cur.close()
            if conn:
                conn.close()

    
    def process_directory(self, directory: pathlib.Path, recursive: bool = False, batch_size=64):
        """Обработка директории с поддержкой рекурсии"""
        all_documents = []
        
        # Поддерживаемые расширения
        patterns = ['*.txt', '*.md', '*.pdf', '*.csv', '*.docx', '*.xlsx']
        
        # Выбор метода поиска
        search_method = directory.rglob if recursive else directory.glob
        
        for pattern in patterns:
            for filepath in search_method(pattern):
                if filepath.stat().st_size == 0:
                    logger.warning(f"Пропускаем пустой файл: {filepath}")
                    continue
                
                size = filepath.stat().st_size
                if filepath.suffix.lower() in (".docx", ".xlsx"):
                    if size > 70 * 1024 * 1024:
                        logger.warning("DOCX/XLSX слишком большой, пропускаем: %s", filepath)
                        continue
                elif filepath.suffix.lower() ==  ".pdf":
                    if size > 200 * 1024 * 1024:
                        logger.warning("PDF слишком большой, пропускаем: %s", filepath)
                        continue
                else:
                    if size > 10 * 1024 * 1024:
                        logger.warning("Файл слишком большой, пропускаем: %s", filepath)
                        continue
                    
                logger.info(f"Обработка {filepath}...")
                try:
                    documents = self.process_txt(filepath)
                    all_documents.extend(documents)
                    logger.info(f"  → Извлечено {len(documents)} фрагментов")
                except Exception as e:
                    logger.error(f"  ✗ Ошибка при обработке {filepath}: {e}")
        
        if all_documents:
            # Сортируем по важности перед сохранением
            all_documents.sort(key=lambda d: d.importance_score, reverse=True)
            self.save_to_db(all_documents, batch_size=batch_size, clean_start=False)
        else:
            logger.warning("Не найдено документов для обработки")
        
        return len(all_documents)


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Улучшенная индексация документов')
    parser.add_argument('--dir', type=str, default='documents', help='Директория с документами')
    # parser.add_argument('--recursive', action='store_true', help='Рекурсивный поиск')
    parser.add_argument('--no-recursive', dest='recursive', action='store_false',
                        help='Отключить рекурсивный поиск')
    parser.set_defaults(recursive=True)
    parser.add_argument('--batch-size', type=int, default=16, help='Размер батча')
    args = parser.parse_args()
    
    # Конфигурация БД
    DB_CONFIG = {
        "dbname": "rag",
        "user": "rag",
        "password": "rag",
        "host": "127.0.0.1",
        "port": 5432
    }
    
    try:
        processor = AdvancedDocumentProcessor(DB_CONFIG)
        
        docs_dir = pathlib.Path(args.dir)
        if docs_dir.exists():
            count = processor.process_directory(
                docs_dir, recursive=args.recursive, batch_size=args.batch_size
            )
            logger.info(f"Обработка завершена. Всего документов: {count}")
        else:
            logger.error(f"Директория '{args.dir}' не существует")
            sys.exit(1)
            
    except Exception as e:
        logger.error(f"Критическая ошибка: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)